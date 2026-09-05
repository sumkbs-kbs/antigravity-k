"""P3-B 테스트: 데이터 레시피 프리셋 (unsloth Data Recipes 벤치마킹).

레시피 카탈로그, 소스 변환(CSV/JSONL/문서), LoRAPipeline 연동을 검증한다.
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

import pytest

from antigravity_k.engine.data_recipes import (
    RECIPES,
    MissingDocumentParserError,
    UnknownRecipeError,
    format_recipe_plan,
    get_recipe,
    list_recipes,
    load_records_from_source,
    normalize_recipe_name,
    records_to_training_jsonl,
)
from antigravity_k.engine.lora_pipeline import LoRAPipeline

try:  # documents extra — 설치된 환경에서만 PDF/DOCX 실변환 테스트 실행
    import docx  # noqa: F401
    import pypdf  # noqa: F401

    DOCUMENTS_EXTRA_INSTALLED = True
except ImportError:
    DOCUMENTS_EXTRA_INSTALLED = False

# ─── 카탈로그 ─────────────────────────────────────────────────────────


def test_catalog_contains_expected_presets() -> None:
    names = {recipe.name for recipe in RECIPES}
    expected = {"chat-sft", "instruction-sft", "preference-dpo", "csv-to-chat", "jsonl-to-chat", "docs-qa-sft"}
    assert expected <= names


def test_catalog_includes_pdf_docx_presets() -> None:
    """PDF/DOCX 프리셋이 카탈로그에 있고 documents extra 필요성을 명시한다."""
    pdf = get_recipe("pdf-qa-sft")
    docx_recipe = get_recipe("docx-qa-sft")
    assert pdf.format == "chat" and docx_recipe.format == "chat"
    assert "pdf" in pdf.source_hint and "docx" in docx_recipe.source_hint
    assert any("extra" in tag or "documents" in tag for tag in (*pdf.tags, *docx_recipe.tags))


def test_list_recipes_is_serializable() -> None:
    listing = list_recipes()
    assert listing
    for entry in listing:
        assert {"name", "title", "format", "min_records"} <= set(entry)


def test_get_recipe_case_insensitive() -> None:
    assert get_recipe("CHAT-SFT").name == "chat-sft"


def test_get_recipe_unknown_raises() -> None:
    with pytest.raises(UnknownRecipeError) as exc_info:
        get_recipe("no-such-recipe")
    assert "no-such-recipe" in str(exc_info.value)
    assert "chat-sft" in str(exc_info.value)  # 지원 목록 안내


def test_normalize_recipe_name() -> None:
    assert normalize_recipe_name("My Custom Recipe!") == "my-custom-recipe"


# ─── 소스 변환 ────────────────────────────────────────────────────────


def test_load_csv_records(tmp_path: Path) -> None:
    csv_file = tmp_path / "data.csv"
    csv_file.write_text(
        "prompt,response\n서울의 수도는?,파리\n이름이 뭔가요?,철수\n",
        encoding="utf-8",
    )
    records = load_records_from_source(str(csv_file))
    assert len(records) == 2
    assert records[0] == {"prompt": "서울의 수도는?", "response": "파리"}


def test_load_csv_skips_incomplete_rows(tmp_path: Path) -> None:
    csv_file = tmp_path / "partial.csv"
    csv_file.write_text(
        "prompt,response\n질문,답변\n답 없는 질문,\n",
        encoding="utf-8",
    )
    records = load_records_from_source(str(csv_file))
    assert len(records) == 1


def test_load_jsonl_alpaca_format(tmp_path: Path) -> None:
    jsonl_file = tmp_path / "alpaca.jsonl"
    jsonl_file.write_text(
        json.dumps({"instruction": "안녕", "output": "반갑다"}, ensure_ascii=False)
        + "\n"
        + json.dumps({"instruction": "누구야", "output": "나야"}, ensure_ascii=False)
        + "\n",
        encoding="utf-8",
    )
    records = load_records_from_source(str(jsonl_file))
    assert len(records) == 2
    assert records[0]["prompt"] == "안녕"


def test_load_jsonl_chatml_format(tmp_path: Path) -> None:
    jsonl_file = tmp_path / "chat.jsonl"
    jsonl_file.write_text(
        json.dumps(
            {"messages": [{"role": "user", "content": "질문"}, {"role": "assistant", "content": "응답"}]},
            ensure_ascii=False,
        )
        + "\n",
        encoding="utf-8",
    )
    records = load_records_from_source(str(jsonl_file))
    assert records == [{"prompt": "질문", "response": "응답"}]


def test_load_md_documents_as_qa(tmp_path: Path) -> None:
    md_file = tmp_path / "guide.md"
    md_file.write_text(
        "# 가이드\n\n## 설치 방법\n\npip install ssak-ai 로 설치합니다. 이 설치 과정은 매우 간단합니다.\n\n## 사용법\n\nuv run agk run 으로 실행합니다. 실행 후 프롬프트에 작업을 입력하세요.\n",
        encoding="utf-8",
    )
    records = load_records_from_source(str(md_file))
    assert len(records) == 2
    assert records[0]["prompt"] == "설치 방법"
    assert "pip install" in records[0]["response"]


def test_load_missing_file_raises() -> None:
    with pytest.raises(FileNotFoundError):
        load_records_from_source("/nonexistent/path/data.csv")


def test_load_unsupported_extension_raises(tmp_path: Path) -> None:
    bad = tmp_path / "data.xlsx"
    bad.write_text("binary", encoding="utf-8")
    with pytest.raises(ValueError, match="지원하지 않는"):
        load_records_from_source(str(bad))


# ─── PDF/DOCX 소스 (documents extra) ─────────────────────────────────


def _write_pdf_with_text(path: Path, page_headers: list[str]) -> None:
    """텍스트 레이어가 있는 실제 PDF 생성 (pypdf writer + content stream 직접 구성).

    각 페이지: 헤더 라인 1개 + 본문 2문장 (Helvetica 기준 latin-1 인코딩 텍스트).
    """
    from pypdf import PdfWriter
    from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

    writer = PdfWriter()

    font = DictionaryObject()
    font[NameObject("/Type")] = NameObject("/Font")
    font[NameObject("/Subtype")] = NameObject("/Type1")
    font[NameObject("/BaseFont")] = NameObject("/Helvetica")
    font_ref = writer._add_object(font)

    resources = DictionaryObject()
    resources[NameObject("/Font")] = DictionaryObject({NameObject("/F1"): font_ref})

    for header in page_headers:
        pdf_lines = [
            header,
            "Run pip install ssak-ai to install the package quickly and easily.",
            "The installer configures all required dependencies automatically for you.",
        ]
        content_parts = [b"BT /F1 12 Tf 72 720 Td"]
        for i, line in enumerate(pdf_lines):
            if i > 0:
                content_parts.append(b" 0 -20 Td")
            escaped = line.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
            content_parts.append(f" ({escaped}) Tj".encode("latin-1", errors="replace"))
        content_parts.append(b" ET")
        stream = DecodedStreamObject()
        stream.set_data(b"".join(content_parts))
        page = writer.add_blank_page(width=612, height=792)
        page[NameObject("/Contents")] = writer._add_object(stream)
        page[NameObject("/Resources")] = resources

    with open(path, "wb") as f:
        writer.write(f)


def test_records_to_training_jsonl_chat(tmp_path: Path) -> None:
    out = tmp_path / "dataset.jsonl"
    count = records_to_training_jsonl(
        [{"prompt": "질문", "response": "답변"}],
        str(out),
        fmt="chat",
    )
    assert count == 1
    record = json.loads(out.read_text(encoding="utf-8"))
    assert record["messages"][0] == {"role": "user", "content": "질문"}
    assert record["messages"][1] == {"role": "assistant", "content": "답변"}


def test_records_to_training_jsonl_instruction(tmp_path: Path) -> None:
    out = tmp_path / "dataset.jsonl"
    records_to_training_jsonl([{"prompt": "질문", "response": "답변"}], str(out), fmt="instruction")
    record = json.loads(out.read_text(encoding="utf-8"))
    assert record["instruction"] == "질문"
    assert record["output"] == "답변"


# ─── LoRAPipeline 연동 ────────────────────────────────────────────────


def test_apply_recipe_with_harvest_source(tmp_path: Path) -> None:
    """수확 소스 경로: export → config 생성 + 레시피 하이퍼파라미터 병합."""
    pipeline = LoRAPipeline(harvest_dir=str(tmp_path / "harvest"))
    result = pipeline.apply_recipe(
        "chat-sft",
        base_model="test-model",
        output_dir=str(tmp_path / "out"),
    )
    assert result["recipe"] == "chat-sft"
    assert result["format"] == "chat"
    assert result["sufficient"] is False  # 수확 데이터 없음
    assert Path(str(result["config_path"])).is_file()
    config = json.loads(Path(str(result["config_path"])).read_text(encoding="utf-8"))
    assert config["recipe"] == "chat-sft"
    assert config["hyperparameters"]["iterations"] == 600  # 레시피 조정 반영


def test_apply_recipe_with_csv_source(tmp_path: Path) -> None:
    csv_file = tmp_path / "qa.csv"
    rows = "\n".join(f"질문{i},답변{i}" for i in range(25))
    csv_file.write_text("prompt,response\n" + rows + "\n", encoding="utf-8")
    pipeline = LoRAPipeline(harvest_dir=str(tmp_path / "harvest"))
    result = pipeline.apply_recipe(
        "csv-to-chat",
        base_model="test-model",
        output_dir=str(tmp_path / "out"),
        source=str(csv_file),
    )
    assert result["records"] == 25
    assert result["sufficient"] is True
    dataset = Path(str(result["dataset_path"]))
    first = json.loads(dataset.read_text(encoding="utf-8").splitlines()[0])
    assert first["messages"][0]["content"] == "질문0"


def test_apply_recipe_dpo_uses_pairs(tmp_path: Path) -> None:
    pipeline = LoRAPipeline(harvest_dir=str(tmp_path / "harvest"))
    _ = pipeline.record_pair("질문", "좋은 답", "나쁜 답", 0.9, 0.3)
    result = pipeline.apply_recipe(
        "preference-dpo",
        base_model="test-model",
        output_dir=str(tmp_path / "out"),
    )
    assert result["records"] == 1
    assert result["format"] == "dpo"


def test_apply_recipe_unknown_raises(tmp_path: Path) -> None:
    pipeline = LoRAPipeline(harvest_dir=str(tmp_path / "harvest"))
    with pytest.raises(UnknownRecipeError):
        pipeline.apply_recipe("mystery", base_model="m", output_dir=str(tmp_path / "out"))


def test_format_recipe_plan_flags_insufficient() -> None:
    recipe = get_recipe("chat-sft")
    plan = format_recipe_plan(recipe, records=3)
    assert "미달" in plan
    assert "chat" in plan


def test_unsupported_extension_message_lists_pdf_docx(tmp_path: Path) -> None:
    """에러 메시지에 pdf/docx 지원 형식이 안내된다."""
    bad = tmp_path / "data.xlsx"
    bad.write_text("binary", encoding="utf-8")
    with pytest.raises(ValueError, match="docx"):
        load_records_from_source(str(bad))


def test_pdf_missing_pypdf_gives_install_hint(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    """pypdf 미설치 환경이면 설치 안내와 함께 MissingDocumentParserError."""
    monkeypatch.setitem(sys.modules, "pypdf", None)  # import → ImportError 시뮬레이션
    pdf_file = tmp_path / "guide.pdf"
    pdf_file.write_bytes(b"%PDF-1.4 fake")

    with pytest.raises(MissingDocumentParserError, match="documents"):
        load_records_from_source(str(pdf_file))


def test_docx_missing_python_docx_gives_install_hint(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setitem(sys.modules, "docx", None)
    docx_file = tmp_path / "manual.docx"
    docx_file.write_bytes(b"PK fake")

    with pytest.raises(MissingDocumentParserError, match="documents"):
        load_records_from_source(str(docx_file))


# ─── PDF/DOCX 소스 (documents extra) ─────────────────────────────────


@pytest.mark.skipif(not DOCUMENTS_EXTRA_INSTALLED, reason="pypdf/python-docx 미설치 (documents extra)")
class TestPdfSource:
    def test_load_pdf_pages_as_qa(self, tmp_path: Path) -> None:
        pdf_file = tmp_path / "guide.pdf"
        _write_pdf_with_text(pdf_file, ["Install Guide", "Usage Basics"])

        records = load_records_from_source(str(pdf_file))

        assert len(records) == 2
        assert records[0]["prompt"] == "Install Guide"
        assert "pip install" in records[0]["response"]
        assert records[1]["prompt"] == "Usage Basics"

    def test_load_pdf_mixed_with_other_sources(self, tmp_path: Path) -> None:
        """PDF + CSV 쉼표 목록 소스가 함께 로드된다."""
        pdf_file = tmp_path / "guide.pdf"
        _write_pdf_with_text(pdf_file, ["Install Guide"])
        csv_file = tmp_path / "qa.csv"
        csv_file.write_text("prompt,response\n질문,답변\n", encoding="utf-8")

        records = load_records_from_source(f"{pdf_file},{csv_file}")

        assert len(records) == 2
        assert records[0]["prompt"] == "Install Guide"
        assert records[1] == {"prompt": "질문", "response": "답변"}

    def test_apply_recipe_pdf_source(self, tmp_path: Path) -> None:
        pdf_file = tmp_path / "guide.pdf"
        _write_pdf_with_text(pdf_file, ["Chapter One", "Chapter Two"])
        pipeline = LoRAPipeline(harvest_dir=str(tmp_path / "harvest"))

        result = pipeline.apply_recipe(
            "pdf-qa-sft",
            base_model="test-model",
            output_dir=str(tmp_path / "out"),
            source=str(pdf_file),
        )

        assert result["records"] == 2
        assert result["recipe"] == "pdf-qa-sft"
        dataset = Path(str(result["dataset_path"]))
        first = json.loads(dataset.read_text(encoding="utf-8").splitlines()[0])
        assert first["messages"][0]["content"] == "Chapter One"


@pytest.mark.skipif(not DOCUMENTS_EXTRA_INSTALLED, reason="pypdf/python-docx 미설치 (documents extra)")
class TestDocxSource:
    def test_load_docx_headings_as_qa(self, tmp_path: Path) -> None:
        import docx

        docx_file = tmp_path / "manual.docx"
        document = docx.Document()
        document.add_heading("설치 방법", level=1)
        document.add_paragraph("pip install ssak-ai 명령으로 설치합니다. 설치 과정은 매우 간단합니다.")
        document.add_heading("사용 방법", level=1)
        document.add_paragraph("uv run agk run 명령으로 실행합니다. 프롬프트에 작업을 입력하세요.")
        document.save(str(docx_file))

        records = load_records_from_source(str(docx_file))

        assert len(records) == 2
        assert records[0]["prompt"] == "설치 방법"
        assert "pip install" in records[0]["response"]
        assert records[1]["prompt"] == "사용 방법"

    def test_load_docx_without_headings_falls_back_to_summary(self, tmp_path: Path) -> None:
        import docx

        docx_file = tmp_path / "plain.docx"
        document = docx.Document()
        document.add_paragraph("헤딩이 없는 일반 문서입니다. " * 5)
        document.save(str(docx_file))

        records = load_records_from_source(str(docx_file))

        assert len(records) == 1
        assert "정리해줘" in records[0]["prompt"]

    def test_apply_recipe_docx_source(self, tmp_path: Path) -> None:
        import docx

        docx_file = tmp_path / "manual.docx"
        document = docx.Document()
        document.add_heading("설치", level=1)
        document.add_paragraph("pip install ssak-ai 명령으로 설치합니다. 설치 후 바로 사용할 수 있습니다.")
        document.save(str(docx_file))
        pipeline = LoRAPipeline(harvest_dir=str(tmp_path / "harvest"))

        result = pipeline.apply_recipe(
            "docx-qa-sft",
            base_model="test-model",
            output_dir=str(tmp_path / "out"),
            source=str(docx_file),
        )

        assert result["records"] == 1
        assert result["recipe"] == "docx-qa-sft"


# ─── Phase 23 P2: mlx 플랫폼 train/valid 자동 분할 테스트 ───────────────


def test_split_dataset_for_mlx_replaces_under_threshold(tmp_path: Path) -> None:
    """데이터가 2 * batch_size 미만인 경우(예: 6건, batch_size=4) train과 valid에 동일 데이터 재사용."""
    src = tmp_path / "raw.jsonl"
    lines = [f'{{"messages": [{{"role": "user", "content": "q{i}"}}]}}\n' for i in range(6)]
    src.write_text("".join(lines), encoding="utf-8")

    out_dir = tmp_path / "mlx_dataset"
    train_path, valid_path = LoRAPipeline.split_dataset_for_mlx(src, out_dir, batch_size=4)

    assert train_path == out_dir / "train.jsonl"
    assert valid_path == out_dir / "valid.jsonl"
    assert (out_dir / "test.jsonl").exists() is False  # 빈 test.jsonl 생성 방지 확인

    train_lines = train_path.read_text(encoding="utf-8").splitlines()
    valid_lines = valid_path.read_text(encoding="utf-8").splitlines()
    assert len(train_lines) == 6
    assert len(valid_lines) == 6
    assert train_lines == [line.strip() for line in lines]
    assert valid_lines == [line.strip() for line in lines]


def test_split_dataset_for_mlx_splits_partition(tmp_path: Path) -> None:
    """데이터가 충분한 경우(25건, batch_size=4) disjoint하게 분할되고 valid에 최소 batch_size 보장."""
    src = tmp_path / "raw.jsonl"
    lines = [f'{{"id": {i}}}\n' for i in range(25)]
    src.write_text("".join(lines), encoding="utf-8")

    out_dir = tmp_path / "mlx_dataset"
    train_path, valid_path = LoRAPipeline.split_dataset_for_mlx(src, out_dir, batch_size=4, train_ratio=0.9, seed=123)

    train_lines = train_path.read_text(encoding="utf-8").splitlines()
    valid_lines = valid_path.read_text(encoding="utf-8").splitlines()

    assert len(train_lines) + len(valid_lines) == 25
    assert len(valid_lines) >= 4
    assert len(train_lines) >= 4
    # 상호 배타적 분할 검증 (disjoint sets)
    assert set(train_lines).isdisjoint(set(valid_lines))
    assert set(train_lines) | set(valid_lines) == {line.strip() for line in lines}


def test_split_dataset_for_mlx_empty(tmp_path: Path) -> None:
    """빈 파일인 경우 빈 train.jsonl / valid.jsonl 생성."""
    src = tmp_path / "empty.jsonl"
    src.write_text("", encoding="utf-8")

    out_dir = tmp_path / "mlx_dataset"
    train_path, valid_path = LoRAPipeline.split_dataset_for_mlx(src, out_dir, batch_size=4)

    assert train_path.read_text(encoding="utf-8") == ""
    assert valid_path.read_text(encoding="utf-8") == ""


def test_apply_recipe_mlx_platform_creates_train_valid_layout(tmp_path: Path) -> None:
    """platform='mlx' 지정 시 apply_recipe가 mlx_dataset/train.jsonl + valid.jsonl로 자동 분할."""
    csv_file = tmp_path / "qa.csv"
    rows = "\n".join(f"질문{i},답변{i}" for i in range(20))
    csv_file.write_text("prompt,response\n" + rows + "\n", encoding="utf-8")

    pipeline = LoRAPipeline(harvest_dir=str(tmp_path / "harvest"))
    out_dir = tmp_path / "out"
    result = pipeline.apply_recipe(
        "csv-to-chat",
        base_model="mlx-community/Qwen2.5-0.5B-4bit",
        output_dir=str(out_dir),
        source=str(csv_file),
        platform="mlx",
        hyperparameter_overrides={"batch_size": 4},
    )

    mlx_dir = out_dir / "mlx_dataset"
    assert result["dataset_path"] == str(mlx_dir)
    assert result["train_path"] == str(mlx_dir / "train.jsonl")
    assert result["valid_path"] == str(mlx_dir / "valid.jsonl")
    assert result["source_dataset_path"] == str(out_dir / "recipe_dataset.jsonl")

    assert (mlx_dir / "train.jsonl").is_file()
    assert (mlx_dir / "valid.jsonl").is_file()
    assert (out_dir / "recipe_dataset.jsonl").is_file()
    assert (mlx_dir / "test.jsonl").exists() is False

    config = result["config"]
    assert config["platform"] == "mlx"
    assert config["dataset"] == str(mlx_dir)
    assert f"--data {mlx_dir}" in str(config["command"])
    assert config["train_path"] == str(mlx_dir / "train.jsonl")
    assert config["valid_path"] == str(mlx_dir / "valid.jsonl")

    # 분할 확인 (20건, batch_size=4 -> train 16, valid 4)
    train_lines = (mlx_dir / "train.jsonl").read_text(encoding="utf-8").splitlines()
    valid_lines = (mlx_dir / "valid.jsonl").read_text(encoding="utf-8").splitlines()
    assert len(train_lines) >= 4
    assert len(valid_lines) >= 4
    assert len(train_lines) + len(valid_lines) == 20


def test_apply_recipe_mlx_platform_with_harvest(tmp_path: Path) -> None:
    """platform='mlx'로 harvest 소스 적용 시 6건 데이터가 재사용 분할되어 생성된다."""
    pipeline = LoRAPipeline(harvest_dir=str(tmp_path / "harvest"))
    for i in range(6):
        pipeline.harvest(f"요청{i}", f"응답{i}", quality_score=0.9)

    out_dir = tmp_path / "out_harvest"
    result = pipeline.apply_recipe(
        "chat-sft",
        base_model="mlx-community/Qwen2.5-0.5B-4bit",
        output_dir=str(out_dir),
        platform="mlx",
    )

    mlx_dir = out_dir / "mlx_dataset"
    assert result["dataset_path"] == str(mlx_dir)
    train_lines = (mlx_dir / "train.jsonl").read_text(encoding="utf-8").splitlines()
    valid_lines = (mlx_dir / "valid.jsonl").read_text(encoding="utf-8").splitlines()
    # 6건 < 2 * 4이므로 동일 6건 재사용
    assert len(train_lines) == 6
    assert len(valid_lines) == 6


def test_apply_recipe_mlx_platform_dpo(tmp_path: Path) -> None:
    """platform='mlx'로 DPO(DoRA) 레시피 적용 시 mlx_dataset이 생성된다."""
    pipeline = LoRAPipeline(harvest_dir=str(tmp_path / "harvest"))
    for i in range(3):
        _ = pipeline.record_pair(f"질문{i}", f"좋은답{i}", f"나쁜답{i}", 0.9, 0.2)

    out_dir = tmp_path / "out_dpo"
    result = pipeline.apply_recipe(
        "preference-dpo",
        base_model="test-model",
        output_dir=str(out_dir),
        platform="mlx",
    )

    mlx_dir = out_dir / "mlx_dataset"
    assert result["dataset_path"] == str(mlx_dir)
    assert (mlx_dir / "train.jsonl").is_file()
    assert (mlx_dir / "valid.jsonl").is_file()
