"""pdf_source_options 테스트 — 페이지 범위·헤더 필터 (Phase 21).

단위 테스트(파서/필터 로직) + 실제 PDF를 사용하는 통합 테스트
(`documents` extra 미설치 환경에서는 통합 테스트가 스킵된다).
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest

from antigravity_k.engine.pdf_source_options import PdfSourceOptions, parse_page_ranges

try:
    import pypdf  # noqa: F401

    DOCUMENTS_EXTRA_INSTALLED = True
except ImportError:
    DOCUMENTS_EXTRA_INSTALLED = False


# ─── 페이지 범위 파서 ────────────────────────────────────────────────


class TestParsePageRanges:
    def test_empty_spec_returns_none_for_all_pages(self) -> None:
        assert parse_page_ranges("", 10) is None
        assert parse_page_ranges(None, 10) is None
        assert parse_page_ranges("  ", 10) is None

    def test_single_page(self) -> None:
        assert parse_page_ranges("3", 10) == [3]

    def test_ranges_and_mixes_are_sorted_and_deduped(self) -> None:
        assert parse_page_ranges("1-5,8,11-13", 20) == [1, 2, 3, 4, 5, 8, 11, 12, 13]
        assert parse_page_ranges("8,1-3,8", 10) == [1, 2, 3, 8]

    def test_boundary_page_allowed(self) -> None:
        assert parse_page_ranges("10", 10) == [10]
        assert parse_page_ranges("9-10", 10) == [9, 10]

    @pytest.mark.parametrize("bad", ["0", "5-2", "abc", "1-", "11", "1-11", "-3", ",,", "1.5"])
    def test_invalid_specs_raise(self, bad: str) -> None:
        with pytest.raises(ValueError, match="페이지"):
            parse_page_ranges(bad, 10)

    def test_whitespace_tolerant(self) -> None:
        assert parse_page_ranges(" 2 - 4 , 7 ", 10) == [2, 3, 4, 7]


# ─── 헤더 필터 ───────────────────────────────────────────────────────


class TestPdfSourceOptionsFilter:
    def test_no_filter_matches_everything(self) -> None:
        opts = PdfSourceOptions()
        assert opts.header_matches("anything") is True
        assert opts.header_matches("") is True
        assert opts.exclude_mode is False

    def test_include_filter(self) -> None:
        opts = PdfSourceOptions(header_filter="Chapter \\d+")
        assert opts.header_matches("Chapter 3: Intro") is True
        assert opts.header_matches("Appendix A") is False
        assert opts.exclude_mode is False

    def test_exclude_filter_with_bang_prefix(self) -> None:
        opts = PdfSourceOptions(header_filter="!TOC")
        assert opts.header_matches("TOC") is False
        assert opts.header_matches("Preface") is True
        assert opts.exclude_mode is True

    def test_search_semantics_not_fullmatch(self) -> None:
        opts = PdfSourceOptions(header_filter="FAQ")
        assert opts.header_matches("Product FAQ list") is True

    def test_invalid_regex_raises_value_error(self) -> None:
        with pytest.raises(ValueError, match="정규식"):
            PdfSourceOptions(header_filter="[unclosed")

    def test_is_immutable(self) -> None:
        opts = PdfSourceOptions(pages="1", header_filter="X")
        with pytest.raises(AttributeError):
            opts.pages = "2"  # type: ignore[misc]


# ─── 질문 템플릿 (Phase 48) ──────────────────────────────────────


class TestQuestionTemplate:
    def test_render_replaces_known_placeholders(self) -> None:
        from antigravity_k.engine.pdf_source_options import render_question_template

        out = render_question_template(
            "{title} 문서의 {page}페이지: {header}에 대해 설명해줘",
            page=3,
            title="manual",
            header="Usage",
            body="본문",
        )
        assert out == "manual 문서의 3페이지: Usage에 대해 설명해줘"

    def test_render_leaves_unknown_tokens_untouched(self) -> None:
        from antigravity_k.engine.pdf_source_options import render_question_template

        # str.format 주입 대비 — 알 수 없는 {토큰}은 그대로 유지 (에러 아님)
        out = render_question_template("{page}쪽 {evil} {0}", page=1, title="t", header="", body="b")
        assert out == "1쪽 {evil} {0}"

    def test_has_question_template_flag(self) -> None:
        assert PdfSourceOptions().has_question_template is False
        assert PdfSourceOptions(question_template="요약해줘").has_question_template is True

    def test_template_is_stripped(self) -> None:
        opts = PdfSourceOptions(question_template="  {page}페이지 정리  ")
        assert opts.question_template == "{page}페이지 정리"


# ─── 실제 PDF 통합 (documents extra 필요) ────────────────────────────


def _write_pdf_with_text(path: Path, page_headers: list[str]) -> None:
    """tests/test_data_recipes.py와 동일한 실제 PDF 생성 방식."""
    from pypdf import PdfWriter
    from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

    writer = PdfWriter()
    font = DictionaryObject()
    font[NameObject("/Type")] = NameObject("/Font")
    font[NameObject("/Subtype")] = NameObject("/Type1")
    font[NameObject("/BaseFont")] = NameObject("/Helvetica")
    font_ref = writer._add_object(font)

    resources = DictionaryObject()
    resources[NameObject("/Font")] = DictionaryObject({NameObject("/F1"): font_ref})

    for header in page_headers:
        pdf_lines = [header, "Body text for this page provides enough content to be useful here."]
        content_parts = [b"BT /F1 12 Tf 72 720 Td"]
        for i, line in enumerate(pdf_lines):
            if i > 0:
                content_parts.append(b" 0 -20 Td")
            escaped = line.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
            content_parts.append(f" ({escaped}) Tj".encode("latin-1", errors="replace"))
        content_parts.append(b" ET")
        stream = DecodedStreamObject()
        stream.set_data(b"".join(content_parts))
        page = writer.add_blank_page(width=612, height=792)
        page[NameObject("/Contents")] = writer._add_object(stream)
        page[NameObject("/Resources")] = resources

    with open(path, "wb") as f:
        writer.write(f)


@pytest.mark.skipif(not DOCUMENTS_EXTRA_INSTALLED, reason="pypdf 미설치 (documents extra)")
class TestPdfRecordsWithOptions:
    def _headers(self) -> list[str]:
        return ["Install Guide", "TOC Page", "Usage Basics", "FAQ List", "Troubleshooting"]

    def test_no_options_returns_all_pages(self, tmp_path: Path) -> None:
        from antigravity_k.engine.data_recipes import _pdf_records

        pdf_file = tmp_path / "doc.pdf"
        _write_pdf_with_text(pdf_file, self._headers())

        records = _pdf_records(pdf_file)
        assert [r["prompt"] for r in records] == self._headers()

    def test_page_range_selects_subset(self, tmp_path: Path) -> None:
        from antigravity_k.engine.data_recipes import _pdf_records

        pdf_file = tmp_path / "doc.pdf"
        _write_pdf_with_text(pdf_file, self._headers())

        records = _pdf_records(pdf_file, options=PdfSourceOptions(pages="2-3"))
        assert [r["prompt"] for r in records] == ["TOC Page", "Usage Basics"]

    def test_page_range_out_of_bounds_raises(self, tmp_path: Path) -> None:
        from antigravity_k.engine.data_recipes import _pdf_records

        pdf_file = tmp_path / "doc.pdf"
        _write_pdf_with_text(pdf_file, self._headers())

        with pytest.raises(ValueError, match="초과"):
            _pdf_records(pdf_file, options=PdfSourceOptions(pages="1-99"))

    def test_include_header_filter_keeps_only_matching(self, tmp_path: Path) -> None:
        from antigravity_k.engine.data_recipes import _pdf_records

        pdf_file = tmp_path / "doc.pdf"
        _write_pdf_with_text(pdf_file, self._headers())

        records = _pdf_records(pdf_file, options=PdfSourceOptions(header_filter="FAQ|Troubleshooting"))
        assert [r["prompt"] for r in records] == ["FAQ List", "Troubleshooting"]

    def test_exclude_header_filter_drops_matching(self, tmp_path: Path) -> None:
        from antigravity_k.engine.data_recipes import _pdf_records

        pdf_file = tmp_path / "doc.pdf"
        _write_pdf_with_text(pdf_file, self._headers())

        records = _pdf_records(pdf_file, options=PdfSourceOptions(header_filter="!TOC"))
        assert [r["prompt"] for r in records] == ["Install Guide", "Usage Basics", "FAQ List", "Troubleshooting"]

    def test_pages_and_filter_combine(self, tmp_path: Path) -> None:
        from antigravity_k.engine.data_recipes import _pdf_records

        pdf_file = tmp_path / "doc.pdf"
        _write_pdf_with_text(pdf_file, self._headers())

        records = _pdf_records(pdf_file, options=PdfSourceOptions(pages="1-4", header_filter="Guide|Basics"))
        assert [r["prompt"] for r in records] == ["Install Guide", "Usage Basics"]

    def test_end_to_end_through_load_records_from_source(self, tmp_path: Path) -> None:
        from antigravity_k.engine.data_recipes import load_records_from_source

        pdf_file = tmp_path / "doc.pdf"
        _write_pdf_with_text(pdf_file, self._headers())

        records = load_records_from_source(
            str(pdf_file), pdf_options=PdfSourceOptions(pages="1,5", header_filter="!TOC")
        )
        assert [r["prompt"] for r in records] == ["Install Guide", "Troubleshooting"]

    def test_question_template_forces_uniform_phrasing(self, tmp_path: Path) -> None:
        """Phase 48: 템플릿 지정 시 헤더가 제목처럼 보여도 템플릿 질문으로 통일된다."""
        from antigravity_k.engine.data_recipes import _pdf_records

        pdf_file = tmp_path / "manual.pdf"
        _write_pdf_with_text(pdf_file, self._headers())

        records = _pdf_records(
            pdf_file,
            options=PdfSourceOptions(question_template="{title}의 {page}페이지 내용을 요약해줘"),
        )
        assert [r["prompt"] for r in records] == [f"manual의 {i}페이지 내용을 요약해줘" for i in (1, 2, 3, 4, 5)]
        # 답변은 기존과 동일하게 페이지 본문 전체
        assert all(len(r["response"]) > 0 for r in records)

    def test_question_template_combines_with_pages_and_filter(self, tmp_path: Path) -> None:
        from antigravity_k.engine.data_recipes import _pdf_records

        pdf_file = tmp_path / "doc.pdf"
        _write_pdf_with_text(pdf_file, self._headers())

        records = _pdf_records(
            pdf_file,
            options=PdfSourceOptions(
                pages="1-4",
                header_filter="Guide|Basics",
                question_template="'{header}' 섹션을 자세히 설명해줘",
            ),
        )
        assert [r["prompt"] for r in records] == [
            "'Install Guide' 섹션을 자세히 설명해줘",
            "'Usage Basics' 섹션을 자세히 설명해줘",
        ]

    def test_question_template_header_placeholder_uses_first_line(self, tmp_path: Path) -> None:
        from antigravity_k.engine.data_recipes import _pdf_records

        pdf_file = tmp_path / "doc.pdf"
        _write_pdf_with_text(pdf_file, self._headers())

        records = _pdf_records(pdf_file, options=PdfSourceOptions(question_template="{header}는 무엇인가?"))
        assert [r["prompt"] for r in records] == [f"{h}는 무엇인가?" for h in self._headers()]

    def test_template_end_to_end_through_load_records_from_source(self, tmp_path: Path) -> None:
        from antigravity_k.engine.data_recipes import load_records_from_source

        pdf_file = tmp_path / "doc.pdf"
        _write_pdf_with_text(pdf_file, self._headers())

        records = load_records_from_source(
            str(pdf_file),
            pdf_options=PdfSourceOptions(pages="2", question_template="{title} {page}페이지 정리"),
        )
        assert [r["prompt"] for r in records] == ["doc 2페이지 정리"]

    def test_end_to_end_through_apply_recipe(self, tmp_path: Path) -> None:
        from antigravity_k.engine.lora_pipeline import LoRAPipeline

        pdf_file = tmp_path / "doc.pdf"
        _write_pdf_with_text(pdf_file, self._headers())
        pipeline = LoRAPipeline(harvest_dir=str(tmp_path / "harvest"))

        result = pipeline.apply_recipe(
            "pdf-qa-sft",
            base_model="test-model",
            output_dir=str(tmp_path / "out"),
            source=str(pdf_file),
            pdf_pages="1-3",
            pdf_header_filter="Guide|TOC|Basics",
        )

        assert result["records"] == 3
        dataset = Path(str(result["dataset_path"]))
        prompts = [
            json.loads(line)["messages"][0]["content"]
            for line in dataset.read_text(encoding="utf-8").splitlines()
            if line.strip()
        ]
        assert prompts == ["Install Guide", "TOC Page", "Usage Basics"]


@pytest.mark.slow
@pytest.mark.skipif(not DOCUMENTS_EXTRA_INSTALLED, reason="pypdf/python-docx 미설치 (documents extra)")
class TestPdfTrainRecipeSlowE2E:
    """요청 시 실행하는 레시피 입력 3경로 통합 E2E (Phase 49 → Phase 50 확장).

    실행: pytest -m slow tests/test_pdf_source_options.py
    PDF/DOCX/CSV 세 소스를 한 번의 패스에 검증:
      - PDF:  pdf-qa-sft  — 페이지 범위 + 헤더 필터 + 질문 템플릿 (선택 옵션 전부)
      - DOCX: docx-qa-sft — 헤딩 섹션 범위 + 질문 템플릿 (Phase 47/48 parity)
      - CSV:  csv-to-chat — 프롬프트/응답 컬럼 직행 (문서 옵션 무시 경로)
    실제 파일 생성 → apply_recipe → 데이터셋/학습 설정 파일까지. 기본 스위트에서는 제외(slow).
    """

    def test_full_train_recipe_writes_dataset_and_config(self, tmp_path: Path) -> None:
        from antigravity_k.engine.lora_pipeline import LoRAPipeline

        headers = ["Install Guide", "TOC Page", "Usage Basics", "FAQ List", "Troubleshooting"]
        pdf_file = tmp_path / "guide.pdf"
        _write_pdf_with_text(pdf_file, headers)
        pipeline = LoRAPipeline(harvest_dir=str(tmp_path / "harvest"))

        result = pipeline.apply_recipe(
            "pdf-qa-sft",
            base_model="mlx-community/Qwen2.5-Coder-32B-Instruct-4bit",
            output_dir=str(tmp_path / "out"),
            source=str(pdf_file),
            pdf_pages="1-5",
            pdf_header_filter="!TOC",
            pdf_question_template="{title} 문서 {page}장: {header}에 대해 설명해줘",
        )

        # 1) 레코드 수: 5페이지 - TOC 1개 제외
        assert result["records"] == 4
        assert result["sufficient"] is False  # 최소 레코드 미달 경고가 포함된 완결 결과

        # 2) 데이터셋 파일 실재 + chat 포맷 + 템플릿 질문
        dataset = Path(str(result["dataset_path"]))
        assert dataset.is_file()
        rows = [json.loads(line) for line in dataset.read_text(encoding="utf-8").splitlines() if line.strip()]
        assert len(rows) == 4
        assert rows[0]["messages"][0]["role"] == "user"
        assert rows[0]["messages"][0]["content"] == "guide 문서 1장: Install Guide에 대해 설명해줘"
        assert "TOC" not in rows[0]["messages"][0]["content"]
        assert all(len(r["messages"][1]["content"]) >= 20 for r in rows)

        # 3) 학습 설정 파일 실재 + pdf-qa-sft 하이퍼파라미터 포함
        config = Path(str(result["config_path"]))
        assert config.is_file()
        config_data = json.loads(config.read_text(encoding="utf-8"))
        assert config_data.get("base_model")
        hyper = config_data.get("hyperparameters") or {}
        assert hyper, "pdf-qa-sft 레시피의 하이퍼파라미터 오버라이드가 설정에 반영돼야 한다"

        # 4) 결과 요약 구조 (CLI가 출력하는 키 전체)
        assert {"recipe", "dataset_path", "config_path", "records", "sufficient"} <= set(result.keys())
        assert result["recipe"] == "pdf-qa-sft"

    # ─── Phase 50: 같은 패스에서 DOCX/CSV 입력 경로까지 ──────────────

    def _assert_recipe_artifacts(self, result: dict, *, recipe: str, expected_rows: int) -> list[dict]:
        """세 소스 공통: 데이터셋/설정 파일 실재 + chat 포맷 행 검증. 파싱된 행 반환."""
        assert result["recipe"] == recipe
        assert result["records"] == expected_rows
        dataset = Path(str(result["dataset_path"]))
        config = Path(str(result["config_path"]))
        assert dataset.is_file() and config.is_file()
        rows = [json.loads(line) for line in dataset.read_text(encoding="utf-8").splitlines() if line.strip()]
        assert len(rows) == expected_rows
        for row in rows:
            assert row["messages"][0]["role"] == "user"
            assert len(row["messages"][1]["content"]) >= 20
        config_data = json.loads(config.read_text(encoding="utf-8"))
        assert config_data.get("base_model")
        assert {"recipe", "dataset_path", "config_path", "records", "sufficient"} <= set(result.keys())
        return rows

    def test_docx_train_recipe_e2e(self, tmp_path: Path) -> None:
        """DOCX 소스: docx-qa-sft가 헤딩 섹션 Q&A 데이터셋/설정을 만든다."""
        import docx

        from antigravity_k.engine.lora_pipeline import LoRAPipeline

        docx_file = tmp_path / "manual.docx"
        document = docx.Document()
        for heading in ["Install Guide", "TOC", "Usage Basics", "FAQ", "Troubleshooting"]:
            document.add_heading(heading, level=1)
            document.add_paragraph(f"{heading} 섹션의 본문 단락입니다. 내용이 충분히 길어야 Q&A로 채택됩니다.")
        document.save(str(docx_file))

        pipeline = LoRAPipeline(harvest_dir=str(tmp_path / "harvest"))
        result = pipeline.apply_recipe(
            "docx-qa-sft",
            base_model="mlx-community/Qwen2.5-Coder-32B-Instruct-4bit",
            output_dir=str(tmp_path / "out-docx"),
            source=str(docx_file),
            pdf_pages="1-4",
            pdf_question_template="{title} 매뉴얼 {page}절 정리",
        )

        rows = self._assert_recipe_artifacts(result, recipe="docx-qa-sft", expected_rows=4)
        questions = [r["messages"][0]["content"] for r in rows]
        assert questions == [f"manual 매뉴얼 {i}절 정리" for i in range(1, 5)]
        assert "TOC" not in " ".join(questions)  # 5번째 섹션(TOC 포함은 2번째) — 범위로 잘림

    def test_csv_train_recipe_e2e(self, tmp_path: Path) -> None:
        """CSV 소스: csv-to-chat가 prompt/response 컬럼을 chat 데이터셋으로 직행한다."""
        import csv as _csv

        from antigravity_k.engine.lora_pipeline import LoRAPipeline

        csv_file = tmp_path / "pairs.csv"
        with open(csv_file, "w", encoding="utf-8", newline="") as f:
            writer = _csv.DictWriter(f, fieldnames=["prompt", "response"])
            writer.writeheader()
            for i in range(1, 6):
                writer.writerow(
                    {"prompt": f"덧셈 {i}+{i}는?", "response": f"{i}+{i}={2 * i} 입니다. 정확히 계산된 답입니다."}
                )

        pipeline = LoRAPipeline(harvest_dir=str(tmp_path / "harvest"))
        result = pipeline.apply_recipe(
            "csv-to-chat",
            base_model="mlx-community/Qwen2.5-Coder-32B-Instruct-4bit",
            output_dir=str(tmp_path / "out-csv"),
            source=str(csv_file),
        )

        rows = self._assert_recipe_artifacts(result, recipe="csv-to-chat", expected_rows=5)
        assert rows[0]["messages"][0]["content"] == "덧셈 1+1는?"
        assert rows[0]["messages"][1]["content"].startswith("1+1=2")

    def test_all_three_sources_yield_disjoint_artifacts(self, tmp_path: Path) -> None:
        """3경로를 각자 출력 디렉터로 연속 실행 — 마지막 실행 아티팩트가 소스와 일치하는지.

        참고: apply_recipe는 dataset 파일명을 고정(recipe_dataset.jsonl)하므로 같은
        output_dir를 공유하면 나중 실행이 덮어쓴다. 호출자가 디렉터를 분리하는 것이 계약 —
        이 테스트는 그 계약대로 분리해서 3경로 연속 실행의 정합성을 검증한다.
        """
        import csv as _csv

        import docx

        from antigravity_k.engine.lora_pipeline import LoRAPipeline

        headers = ["Install Guide", "TOC", "Usage Basics"]
        pdf_file = tmp_path / "a.pdf"
        _write_pdf_with_text(pdf_file, headers)
        docx_file = tmp_path / "b.docx"
        document = docx.Document()
        for heading in ["Alpha", "Beta", "Gamma"]:
            document.add_heading(heading, level=1)
            document.add_paragraph(f"{heading} 섹션 본문. 충분히 긴 단락으로 Q&A 채택 조건을 만족한다.")
        document.save(str(docx_file))
        csv_file = tmp_path / "c.csv"
        with open(csv_file, "w", encoding="utf-8", newline="") as f:
            writer = _csv.DictWriter(f, fieldnames=["prompt", "response"])
            writer.writeheader()
            writer.writerow({"prompt": "질문1", "response": "답변1입니다. 충분히 긴 답변으로 채택 조건을 만족한다."})

        pipeline = LoRAPipeline(harvest_dir=str(tmp_path / "harvest"))
        base = {"base_model": "m"}
        r_pdf = pipeline.apply_recipe("pdf-qa-sft", source=str(pdf_file), output_dir=str(tmp_path / "out-pdf"), **base)
        r_docx = pipeline.apply_recipe(
            "docx-qa-sft", source=str(docx_file), output_dir=str(tmp_path / "out-docx"), **base
        )
        r_csv = pipeline.apply_recipe("csv-to-chat", source=str(csv_file), output_dir=str(tmp_path / "out-csv"), **base)

        # 세 결과의 데이터셋 경로는 모두 다르고(디렉터 분리), 파일이 모두 실재한다.
        paths = {r_pdf["dataset_path"], r_docx["dataset_path"], r_csv["dataset_path"]}
        assert len(paths) == 3
        assert all(Path(p).is_file() for p in paths)
        assert {r_pdf["recipe"], r_docx["recipe"], r_csv["recipe"]} == {"pdf-qa-sft", "docx-qa-sft", "csv-to-chat"}

        # 각 아티팩트 내용이 소스와 일치 — 연속 실행해도 서로 섞이지 않는다.
        def _questions(path: str) -> list[str]:
            lines = Path(path).read_text(encoding="utf-8").splitlines()
            return [json.loads(line)["messages"][0]["content"] for line in lines if line.strip()]

        assert _questions(r_pdf["dataset_path"]) == headers
        assert _questions(r_docx["dataset_path"]) == ["Alpha", "Beta", "Gamma"]
        assert _questions(r_csv["dataset_path"]) == ["질문1"]
