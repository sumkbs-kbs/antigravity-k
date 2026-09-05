"""DOCX 변환기 선택 옵션 테스트 (Phase 47) — pdf-qa-sft와의 parity.

PDF(Phase 21)에 추가했던 페이지 범위/헤더 필터를 DOCX 헤딩 섹션에 동일 적용:
  - `pages`: 헤딩 섹션 번호(1-based, 문서 순서) 선택 — "2-3,5" 문법 공유
  - `header_filter`: 헤딩 텍스트 정규식 포함/"!"접두사 제외 필터
  - 둘 다 PdfSourceOptions 재사용 → 해석 규칙이 PDF와 한곳에 유지된다
"""

from __future__ import annotations

from pathlib import Path

import pytest

from antigravity_k.engine.pdf_source_options import PdfSourceOptions

try:
    import docx  # noqa: F401

    DOCUMENTS_EXTRA_INSTALLED = True
except ImportError:  # pragma: no cover - documents extra 미설치 환경
    DOCUMENTS_EXTRA_INSTALLED = False


def _write_docx_with_sections(path: Path, headings: list[str]) -> None:
    """헤딩 섹션 N개짜리 실제 DOCX 생성. 각 섹션은 충분히 긴 본문 단락 1개를 갖는다."""
    import docx

    document = docx.Document()
    for heading in headings:
        document.add_heading(heading, level=1)
        document.add_paragraph(f"{heading} 섹션의 본문 단락입니다. 내용이 충분히 길어야 Q&A로 채택됩니다.")
    document.save(str(path))


@pytest.mark.skipif(not DOCUMENTS_EXTRA_INSTALLED, reason="python-docx 미설치 (documents extra)")
class TestDocxRecordsWithOptions:
    def _headings(self) -> list[str]:
        return ["Install Guide", "TOC", "Usage Basics", "FAQ", "Troubleshooting"]

    def test_no_options_returns_all_sections(self, tmp_path: Path) -> None:
        from antigravity_k.engine.data_recipes import _docx_records

        docx_file = tmp_path / "doc.docx"
        _write_docx_with_sections(docx_file, self._headings())

        records = _docx_records(docx_file)
        assert [r["prompt"] for r in records] == self._headings()

    def test_section_range_selects_subset(self, tmp_path: Path) -> None:
        from antigravity_k.engine.data_recipes import _docx_records

        docx_file = tmp_path / "doc.docx"
        _write_docx_with_sections(docx_file, self._headings())

        # pages 옵션이 헤딩 섹션 번호 선택으로 동작 (PDF 페이지 범위와 같은 문법)
        records = _docx_records(docx_file, options=PdfSourceOptions(pages="2-3"))
        assert [r["prompt"] for r in records] == ["TOC", "Usage Basics"]

    def test_section_range_out_of_bounds_raises(self, tmp_path: Path) -> None:
        from antigravity_k.engine.data_recipes import _docx_records

        docx_file = tmp_path / "doc.docx"
        _write_docx_with_sections(docx_file, self._headings())

        with pytest.raises(ValueError, match="초과"):
            _docx_records(docx_file, options=PdfSourceOptions(pages="1-99"))

    def test_include_heading_filter_keeps_only_matching(self, tmp_path: Path) -> None:
        from antigravity_k.engine.data_recipes import _docx_records

        docx_file = tmp_path / "doc.docx"
        _write_docx_with_sections(docx_file, self._headings())

        records = _docx_records(docx_file, options=PdfSourceOptions(header_filter="FAQ|Troubleshooting"))
        assert [r["prompt"] for r in records] == ["FAQ", "Troubleshooting"]

    def test_exclude_heading_filter_drops_matching(self, tmp_path: Path) -> None:
        from antigravity_k.engine.data_recipes import _docx_records

        docx_file = tmp_path / "doc.docx"
        _write_docx_with_sections(docx_file, self._headings())

        records = _docx_records(docx_file, options=PdfSourceOptions(header_filter="!TOC"))
        assert [r["prompt"] for r in records] == ["Install Guide", "Usage Basics", "FAQ", "Troubleshooting"]

    def test_range_and_filter_combine(self, tmp_path: Path) -> None:
        from antigravity_k.engine.data_recipes import _docx_records

        docx_file = tmp_path / "doc.docx"
        _write_docx_with_sections(docx_file, self._headings())

        records = _docx_records(docx_file, options=PdfSourceOptions(pages="1-4", header_filter="Guide|Basics"))
        assert [r["prompt"] for r in records] == ["Install Guide", "Usage Basics"]

    def test_no_fallback_summary_when_filter_active(self, tmp_path: Path) -> None:
        """필터가 활성됐는데 매칭이 없으면 빈 결과 — 문서 전체 요약 폴백은 만들지 않는다."""
        from antigravity_k.engine.data_recipes import _docx_records

        docx_file = tmp_path / "doc.docx"
        _write_docx_with_sections(docx_file, self._headings())

        records = _docx_records(docx_file, options=PdfSourceOptions(header_filter="존재하지않는헤딩"))
        assert records == []

    def test_fallback_summary_still_works_without_options(self, tmp_path: Path) -> None:
        """옵션 없 + 헤딩 없는 문서 → 기존대로 문서 전체 요약 Q&A 1건."""
        import docx

        from antigravity_k.engine.data_recipes import _docx_records

        docx_file = tmp_path / "plain.docx"
        plain = docx.Document()
        for i in range(1, 6):
            plain.add_paragraph(f"헤딩 없는 일반 단락 {i}. 문서 전체 요약 폴백이 동작하려면 본문이 충분해야 한다.")
        plain.save(str(docx_file))

        records = _docx_records(docx_file)
        assert len(records) == 1
        assert "plain" in records[0]["prompt"]

    def test_end_to_end_through_load_records_from_source(self, tmp_path: Path) -> None:
        from antigravity_k.engine.data_recipes import load_records_from_source

        docx_file = tmp_path / "doc.docx"
        _write_docx_with_sections(docx_file, self._headings())

        records = load_records_from_source(
            str(docx_file), pdf_options=PdfSourceOptions(pages="1,5", header_filter="!TOC")
        )
        assert [r["prompt"] for r in records] == ["Install Guide", "Troubleshooting"]

    def test_question_template_forces_uniform_phrasing(self, tmp_path: Path) -> None:
        """Phase 48: DOCX도 같은 템플릿 옵션을 받는다 (단위는 헤딩 섹션)."""
        from antigravity_k.engine.data_recipes import _docx_records

        docx_file = tmp_path / "doc.docx"
        _write_docx_with_sections(docx_file, self._headings())

        records = _docx_records(
            docx_file,
            options=PdfSourceOptions(question_template="{title} 문서 {page}번 섹션을 설명해줘"),
        )
        assert [r["prompt"] for r in records] == [f"doc 문서 {i}번 섹션을 설명해줘" for i in range(1, 6)]

    def test_question_template_header_placeholder_is_heading_text(self, tmp_path: Path) -> None:
        from antigravity_k.engine.data_recipes import _docx_records

        docx_file = tmp_path / "doc.docx"
        _write_docx_with_sections(docx_file, self._headings())

        records = _docx_records(
            docx_file,
            options=PdfSourceOptions(pages="2-3", question_template="'{header}'는 무엇을 다루는가?"),
        )
        assert [r["prompt"] for r in records] == ["'TOC'는 무엇을 다루는가?", "'Usage Basics'는 무엇을 다루는가?"]

    def test_docx_and_pdf_accept_the_same_options_object(self, tmp_path: Path) -> None:
        """parity 증명: 같은 PdfSourceOptions 인스턴스가 PDF/DOCX 양쪽에 그대로 전달되고
        동일한 단위(페이지/섹션) 선택 + 헤더/헤딩 필터 의미로 해석된다."""
        from antigravity_k.engine.data_recipes import _docx_records, _pdf_records

        # 3번째 단위(Usage Basics) 선택 + TOC 제외 — 양쪽에서 같은 결과여야 한다.
        shared = PdfSourceOptions(pages="3", header_filter="!TOC")

        docx_file = tmp_path / "doc.docx"
        _write_docx_with_sections(docx_file, self._headings())
        docx_prompts = [r["prompt"] for r in _docx_records(docx_file, options=shared)]

        pdf_prompts: list[str] = []
        try:
            from pypdf import PdfWriter  # noqa: F401

            from tests.test_pdf_source_options import _write_pdf_with_text

            pdf_file = tmp_path / "doc.pdf"
            _write_pdf_with_text(pdf_file, self._headings())
            pdf_prompts = [r["prompt"] for r in _pdf_records(pdf_file, options=shared)]
        except ImportError:
            pytest.skip("pypdf 미설치 — DOCX 단독 검증")

        assert docx_prompts == pdf_prompts == ["Usage Basics"]
