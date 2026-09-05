"""데이터 레시피 (Data Recipes) — 학습 데이터셋 구성 프리셋.

============================================================
벤치마킹 출처: unsloth "Data Recipes" — "Build datasets from PDFs, CSVs, DOCX
files, and more" (no-code 학습 데이터셋 구성).

Ssak-Ai에는 이미 자동 수확(harvest) 기반 SFT/DPO 파이프라인이 있지만,
**어떤 소스를 어떤 학습 포맷으로 어떤 하이퍼파라미터로 학습할지**를 결정하는
프리셋 레이어가 없다. 이 모듈은 결정론적 레시피 카탈로그를 제공한다:

- 레시피 = 소스 유형 + 대상 포맷(SFT/DPO/chat/instruction) + 하이퍼파라미터 조정
- LoRAPipeline에 연결: export_dataset/export_dpo_dataset/generate_config 호출 전 설정 프리셋
- 파일 소스(CSV/JSONL/TXT/MD/PDF/DOCX) → 레시피에 맞는 학습 레코드 변환
- PDF/DOCX 파서(pypdf/python-docx)는 선택 의존성 — `documents` extra로 설치

모든 프리셋은 결정론적(입력이 같으면 결과가 같음)이며, 네트워크 I/O가 없다.
"""

from __future__ import annotations

import csv
import json
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Final, Literal

from antigravity_k.engine.pdf_source_options import PdfSourceOptions, render_question_template

logger = logging.getLogger("antigravity_k.engine.data_recipes")

RecipeFormat = Literal["chat", "instruction", "dpo"]

_INVALID_CHARS: Final[re.Pattern[str]] = re.compile(r"[^a-z0-9_-]")


@dataclass(frozen=True, slots=True)
class DataRecipe:
    """학습 데이터셋 구성 프리셋 1건."""

    name: str  # 정규화된 식별자 (예: "chat-sft")
    title: str
    description: str
    source_hint: str  # 예상 소스 설명 (예: "quality_gate_harvest", "csv:prompt,response")
    format: RecipeFormat
    # 하이퍼파라미터 조정 (generate_config 결과에 병합됨)
    hyperparameter_overrides: dict[str, float | int | str] = field(default_factory=dict)
    # 소스 파일 기반 구성 시 권장 최소 레코드 수
    min_records: int = 10
    tags: tuple[str, ...] = ()


RECIPES: Final[tuple[DataRecipe, ...]] = (
    DataRecipe(
        name="chat-sft",
        title="Chat SFT (수확 데이터)",
        description="QualityGate 수확 데이터를 ChatML로 SFT 학습. 자가 개선 기본 경로.",
        source_hint="quality_gate_harvest",
        format="chat",
        # mlx-lm 0.31.x 기본값(iters 1000, lr 1e-5, batch 4) 기준 조정 —
        # 문서: github.com/ml-explore/mlx-lm/blob/main/mlx_lm/LORA.md
        hyperparameter_overrides={"iterations": 600, "learning_rate": "1e-5", "batch_size": 4},
        min_records=10,
        tags=("sft", "harvest", "chat"),
    ),
    DataRecipe(
        name="instruction-sft",
        title="Alpaca SFT (수확 데이터)",
        description="수확 데이터를 instruction/input/output 포맷으로 SFT. 레거시 호환.",
        source_hint="quality_gate_harvest",
        format="instruction",
        hyperparameter_overrides={"iterations": 600, "learning_rate": "1e-5", "batch_size": 4},
        min_records=10,
        tags=("sft", "harvest", "instruction"),
    ),
    DataRecipe(
        name="preference-dpo",
        title="DPO 선호 정렬",
        description="선호쌍(chosen/rejected)으로 DPO 정렬. SFT 이후 2단계 훈련 권장.",
        source_hint="quality_gate_pairs",
        format="dpo",
        # unsloth 가이드: DPO/RL 계열 학습률 5e-6 (SFT 2e-4 대비 작게), epochs 1-3
        hyperparameter_overrides={
            "iterations": 400,
            "learning_rate": "5e-6",
            "batch_size": 4,
            "num_train_epochs": 1,
        },
        min_records=5,
        tags=("dpo", "preference", "alignment"),
    ),
    DataRecipe(
        name="csv-to-chat",
        title="CSV → Chat SFT",
        description="CSV(prompt,response 필수)를 ChatML 학습 데이터로 변환해 SFT.",
        source_hint="csv:prompt,response",
        format="chat",
        hyperparameter_overrides={"iterations": 800, "learning_rate": "1e-5", "batch_size": 4},
        min_records=20,
        tags=("csv", "sft", "conversion"),
    ),
    DataRecipe(
        name="jsonl-to-chat",
        title="JSONL → Chat SFT",
        description="JSONL(instruction/output 또는 messages)을 ChatML로 정규화해 SFT.",
        source_hint="jsonl:instruction,output|messages",
        format="chat",
        hyperparameter_overrides={"iterations": 800, "learning_rate": "1e-5", "batch_size": 4},
        min_records=20,
        tags=("jsonl", "sft", "conversion"),
    ),
    DataRecipe(
        name="docs-qa-sft",
        title="문서 Q&A → SFT",
        description="TXT/MD 문서를 Q&A 쌍으로 변환(제목→질문, 본문→답)해 SFT.",
        source_hint="files:txt,md",
        format="chat",
        hyperparameter_overrides={"iterations": 500, "learning_rate": "2e-5", "batch_size": 2},
        min_records=5,
        tags=("documents", "qa", "sft"),
    ),
    DataRecipe(
        name="pdf-qa-sft",
        title="PDF → SFT",
        description="PDF를 페이지별 Q&A 쌍으로 변환해 SFT. pypdf 필요 (`documents` extra).",
        source_hint="files:pdf",
        format="chat",
        hyperparameter_overrides={"iterations": 500, "learning_rate": "2e-5", "batch_size": 2},
        min_records=5,
        tags=("pdf", "documents", "qa", "sft"),
    ),
    DataRecipe(
        name="docx-qa-sft",
        title="DOCX → SFT",
        description="DOCX 헤딩 구조를 Q&A 쌍으로 변환해 SFT. python-docx 필요 (`documents` extra).",
        source_hint="files:docx",
        format="chat",
        hyperparameter_overrides={"iterations": 500, "learning_rate": "2e-5", "batch_size": 2},
        min_records=5,
        tags=("docx", "documents", "qa", "sft"),
    ),
)

_RECIPE_INDEX: Final[dict[str, DataRecipe]] = {recipe.name: recipe for recipe in RECIPES}


class UnknownRecipeError(ValueError):
    """존재하지 않는 레시피 이름."""


class MissingDocumentParserError(ValueError):
    """PDF/DOCX 파싱에 필요한 선택 의존성 미설치.

    설치: uv sync --extra documents (또는 pip install 'antigravity-k[documents]')
    """


def list_recipes() -> list[dict[str, object]]:
    """레시피 카탈로그를 UI/CLI용 dict 목록으로 반환."""
    return [
        {
            "name": recipe.name,
            "title": recipe.title,
            "description": recipe.description,
            "source_hint": recipe.source_hint,
            "format": recipe.format,
            "min_records": recipe.min_records,
            "tags": list(recipe.tags),
        }
        for recipe in RECIPES
    ]


def get_recipe(name: str) -> DataRecipe:
    """이름으로 레시피 조회. 없으면 UnknownRecipeError."""
    recipe = _RECIPE_INDEX.get(name.strip().casefold())
    if recipe is None:
        supported = ", ".join(sorted(_RECIPE_INDEX))
        raise UnknownRecipeError(f"알 수 없는 레시피: '{name}'. 지원 목록: {supported}")
    return recipe


def normalize_recipe_name(title: str) -> str:
    """레시피 제목을 정규화된 식별자로 변환."""
    slug = title.strip().casefold().replace(" ", "-")
    return _INVALID_CHARS.sub("", slug)[:60] or "custom-recipe"


# ─── 소스 파일 변환 ──────────────────────────────────────────────────


def _csv_records(path: Path) -> list[dict[str, str]]:
    with open(path, encoding="utf-8", newline="") as f:
        rows = list(csv.DictReader(f))
    result: list[dict[str, str]] = []
    for row in rows:
        prompt = (row.get("prompt") or row.get("instruction") or row.get("question") or "").strip()
        response = (row.get("response") or row.get("output") or row.get("answer") or "").strip()
        if prompt and response:
            result.append({"prompt": prompt, "response": response})
    return result


def _jsonl_records(path: Path) -> list[dict[str, str]]:
    result: list[dict[str, str]] = []
    with open(path, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                raw = json.loads(line)
            except json.JSONDecodeError:
                logger.debug("JSONL 파싱 실패 행 스킵: %s", line[:80])
                continue
            if not isinstance(raw, dict):
                continue
            # 1. ChatML messages 형식
            messages = raw.get("messages")
            if isinstance(messages, list) and messages:
                user = next(
                    (str(m.get("content", "")) for m in messages if isinstance(m, dict) and m.get("role") == "user"),
                    "",
                )
                assistant = next(
                    (
                        str(m.get("content", ""))
                        for m in messages
                        if isinstance(m, dict) and m.get("role") == "assistant"
                    ),
                    "",
                )
                if user and assistant:
                    result.append({"prompt": user, "response": assistant})
                    continue
            # 2. Alpaca 형식
            prompt = str(raw.get("instruction") or raw.get("prompt") or "").strip()
            response = str(raw.get("output") or raw.get("response") or "").strip()
            if prompt and response:
                result.append({"prompt": prompt, "response": response})
    return result


def _docs_records(path: Path) -> list[dict[str, str]]:
    """TXT/MD 문서를 Q&A 쌍으로 변환. 마크다운 헤더(##/###)가 질문, 뒤 본문이 답."""
    pairs: list[dict[str, str]] = []
    if path.suffix.casefold() == ".md":
        text = path.read_text(encoding="utf-8")
        sections = re.split(r"^##\s+", text, flags=re.MULTILINE)
        for section in sections[1:]:
            lines = section.strip().splitlines()
            if not lines:
                continue
            question = lines[0].strip().lstrip("#").strip()
            answer = "\n".join(lines[1:]).strip()
            if question and len(answer) >= 20:
                pairs.append({"prompt": question, "response": answer})
        if pairs:
            return pairs
    # 일반 텍스트/MD 폴백: 단락별로 "이 문서의 요약" 형태 Q&A 1건
    body = path.read_text(encoding="utf-8").strip()
    if len(body) >= 40:
        pairs.append({"prompt": f"{path.stem} 문서의 내용을 정리해줘", "response": body[:4000]})
    return pairs


def _pdf_records(
    path: Path,
    options: PdfSourceOptions | None = None,
) -> list[dict[str, str]]:
    """PDF를 페이지별 Q&A 쌍으로 변환 (unsloth Data Recipes 벤치마킹).

    페이지별 헤더 라인이 있으면 질문으로, 없으면 페이지 요약 질문을 생성한다.
    options로 페이지 범위 선택(`pages`)과 헤더 필터(`header_filter`)를 지정할 수 있다.
    pypdf는 선택 의존성(`documents` extra) — 미설치 시 설치 안내와 함께 실패한다.
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise MissingDocumentParserError(
            "PDF 소스를 읽으려면 pypdf가 필요합니다. 설치: uv sync --extra documents",
        ) from exc

    opts = options or PdfSourceOptions()
    reader = PdfReader(str(path))
    selected_pages = opts.select_pages(len(reader.pages))
    pairs: list[dict[str, str]] = []
    for index, page in enumerate(reader.pages, start=1):
        if selected_pages is not None and index not in selected_pages:
            continue
        text = (page.extract_text() or "").strip()
        if len(text) < 40:
            continue
        lines = text.splitlines()
        # 첫 줄이 짧은 제목처럼 보이면(80자 미만, 마침표 없음) 질문으로 사용
        header = lines[0].strip() if lines else ""
        if header and not opts.header_matches(header):
            continue
        if opts.has_question_template:
            # 템플릿 강제 모드 (Phase 48) — 헤더가 제목처럼 보여도 템플릿 질문을 쓴다.
            body = text
            question = render_question_template(
                opts.question_template,
                page=index,
                title=path.stem,
                header=header,
                body=text,
            )
        elif header and len(header) < 80 and not header.endswith((".", "다.", ":")):
            question = header
            body = "\n".join(lines[1:]).strip() or text
        else:
            question = f"{path.stem} 문서의 {index}페이지 내용을 정리해줘"
            body = text
        pairs.append({"prompt": question, "response": body[:4000]})
    return pairs


def _docx_records(
    path: Path,
    options: PdfSourceOptions | None = None,
) -> list[dict[str, str]]:  # noqa: C901 — PDF parity 옵션 분기라 구조가 단순해도 분기 수가 많다
    """DOCX의 헤딩(Heading 1/2)→질문, 뒤 단락→답 형태로 Q&A 쌍 변환.

    options로 PDF 변환기와 동일한 선택을 지원한다 (Phase 47, pdf-qa-sft와 parity):
      - `pages`: 헤딩 섹션 번호(1-based, 문서 순서) 선택 — "2-3,5" 문법 공유.
      - `header_filter`: 헤딩 텍스트 정규식 포함/"!"접두사 제외 필터.
      둘 다 PDF와 같은 PdfSourceOptions를 재사용하므로 해석 규칙이 한곳에 있다.
    필터가 활성된 문서에서 아무 Q&A도 남지 않으면 폴백 요약은 생성하지 않는다
    (필터 의도를 훼손하지 않기 위해 — PDF 헤더 필터와 동일 원칙).

    python-docx는 선택 의존성(`documents` extra; rag extra와 공유) — 미설치 시
    설치 안내와 함께 실패한다.
    """
    try:
        import docx
    except ImportError as exc:
        raise MissingDocumentParserError(
            "DOCX 소스를 읽으려면 python-docx가 필요합니다. 설치: uv sync --extra documents",
        ) from exc

    from antigravity_k.engine.pdf_source_options import PdfSourceOptions

    opts = options or PdfSourceOptions()
    document = docx.Document(str(path))
    pairs: list[dict[str, str]] = []
    question = ""
    body_parts: list[str] = []
    section_index = 0

    # 섹션 총수 선계산 — 페이지 범위 상한 검증(parse_page_ranges의 total_pages 역할).
    total_sections = sum(
        1
        for p in document.paragraphs
        if p.text.strip()
        and p.style is not None
        and ("heading" in (p.style.name or "").casefold() or "제목" in (p.style.name or ""))
    )
    selected_sections = opts.select_pages(total_sections)

    def _flush() -> None:
        nonlocal question, body_parts
        body = "\n".join(body_parts).strip()
        if question and len(body) >= 20:
            pairs.append({"prompt": question, "response": body[:4000]})
        question = ""
        body_parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "").casefold() if paragraph.style is not None else ""
        is_heading = "heading" in style_name or "제목" in style_name
        if not is_heading:
            if question:
                body_parts.append(text)
            continue
        _flush()
        section_index += 1
        # 섹션 범위/헤딩 필터 둘 다 통과한 헤딩만 질문이 된다 (PDF 페이지+헤더 필터와 동일 조합).
        if selected_sections is not None and section_index not in selected_sections:
            continue
        if not opts.header_matches(text):
            continue
        question = (
            render_question_template(
                opts.question_template,
                page=section_index,
                title=path.stem,
                header=text,
                body="",
            )
            if opts.has_question_template
            else text
        )
    _flush()

    if not pairs and selected_sections is None and not opts.header_filter:
        # 헤딩 없는 문서 폴백: 문서 전체 요약 Q&A 1건.
        # 선택 옵션이 활성된 경우엔 폴백을 만들지 않는다 — "2-3페이지만" 요청에
        # 문서 전체 요약이 섞여 들어가면 PDF 필터와 동일하게 의도 훼손.
        body = "\n".join(p.text.strip() for p in document.paragraphs if p.text.strip())
        if len(body) >= 40:
            pairs.append({"prompt": f"{path.stem} 문서의 내용을 정리해줘", "response": body[:4000]})
    return pairs


def load_records_from_source(
    source: str,
    pdf_options: PdfSourceOptions | None = None,
) -> list[dict[str, str]]:
    """소스(CSV/JSONL/TXT/MD 파일 또는 쉼표 목록)에서 학습 레코드를 로드.

    Args:
        source: 파일 경로(쉼표 목록)
        pdf_options: 문서 소스 옵션 (페이지/섹션 범위, 헤더·헤딩 필터, 질문 템플릿).
            PDF와 DOCX에 적용되며 그 외 형식엔 무시됨 (Phase 47부터 DOCX도 동일 옵션 지원).

    Returns:
        [{"prompt": ..., "response": ...}] 목록

    Raises:
        FileNotFoundError: 파일이 존재하지 않음
        ValueError: 지원하지 않는 확장자, 잘못된 PDF 옵션
        MissingDocumentParserError: PDF/DOCX인데 pypdf/python-docx 미설치

    """
    records: list[dict[str, str]] = []
    seen_paths: set[str] = set()
    for raw_path in source.split(","):
        path = Path(raw_path.strip()).expanduser()
        if not path.is_file():
            raise FileNotFoundError(f"소스 파일을 찾을 수 없습니다: {path}")
        resolved = str(path.resolve())
        if resolved in seen_paths:
            continue
        seen_paths.add(resolved)
        suffix = path.suffix.casefold()
        if suffix == ".csv":
            records.extend(_csv_records(path))
        elif suffix in {".jsonl", ".json"}:
            records.extend(_jsonl_records(path))
        elif suffix in {".txt", ".md"}:
            records.extend(_docs_records(path))
        elif suffix == ".pdf":
            records.extend(_pdf_records(path, options=pdf_options))
        elif suffix in {".docx", ".doc"}:
            records.extend(_docx_records(path, options=pdf_options))
        else:
            raise ValueError(
                f"지원하지 않는 소스 형식: {path.suffix} (csv/jsonl/json/txt/md/pdf/docx 지원)",
            )
    return records


def records_to_training_jsonl(
    records: list[dict[str, str]],
    output_path: str,
    fmt: RecipeFormat = "chat",
) -> int:
    """레코드를 학습용 JSONL로 변환 저장. 반환값은 기록된 행 수."""
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    count = 0
    with open(output, "w", encoding="utf-8") as f:
        for record in records:
            if fmt == "chat":
                payload: dict[str, object] = {
                    "messages": [
                        {"role": "user", "content": record["prompt"]},
                        {"role": "assistant", "content": record["response"]},
                    ],
                }
            else:
                payload = {"instruction": record["prompt"], "input": "", "output": record["response"]}
            _ = f.write(json.dumps(payload, ensure_ascii=False) + "\n")
            count += 1
    logger.info("[DataRecipe] %s건 → %s (%s 포맷)", count, output, fmt)
    return count


def format_recipe_plan(recipe: DataRecipe, records: int, output_path: str = "") -> str:
    """레시피 실행 안내를 마크다운으로 렌더링."""
    lines = [
        f"# {recipe.title}",
        "",
        recipe.description,
        "",
        f"- 포맷: `{recipe.format}`",
        f"- 레코드 수: {records} (권장 최소 {recipe.min_records})",
    ]
    if records < recipe.min_records:
        lines.append("- ⚠️ 권장 최소 레코드 수 미달 — 품질 저하 가능")
    if recipe.hyperparameter_overrides:
        lines.append("- 하이퍼파라미터 조정:")
        lines.extend(f"  - {key}: {value}" for key, value in sorted(recipe.hyperparameter_overrides.items()))
    if output_path:
        lines.append(f"- 데이터셋: `{output_path}`")
    return "\n".join(lines)
